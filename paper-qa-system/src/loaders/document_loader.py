"""
文档加载器模块

支持多种文档格式的加载、文本清洗、元数据提取和错误处理
"""

import os
import re
from pathlib import Path
from typing import List, Optional, Dict, Any, Union
from datetime import datetime

from llama_index.core import Document, SimpleDirectoryReader
from llama_index.core.schema import TextNode
from loguru import logger

# PDF 处理库
try:
    import pymupdf  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF 未安装，将使用 pypdf 作为备选")

try:
    import pypdf
    PYPDF_AVAILABLE = True
except ImportError:
    PYPDF_AVAILABLE = False
    logger.warning("pypdf 未安装")

# DOCX 处理
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装")


class PDFReader:
    """自定义 PDF 读取器，优先使用 PyMuPDF"""
    
    def __init__(self):
        self.use_pymupdf = PYMUPDF_AVAILABLE
        self.use_pypdf = PYPDF_AVAILABLE
        
        if not self.use_pymupdf and not self.use_pypdf:
            raise ImportError("请至少安装 pymupdf 或 pypdf: pip install pymupdf 或 pip install pypdf")
    
    def load_data(self, file_path: Path) -> List[Document]:
        """
        加载 PDF 文件
        
        Args:
            file_path: PDF 文件路径
            
        Returns:
            Document 列表
        """
        try:
            if self.use_pymupdf:
                return self._load_with_pymupdf(file_path)
            elif self.use_pypdf:
                return self._load_with_pypdf(file_path)
        except Exception as e:
            logger.error(f"加载 PDF 失败 {file_path}: {e}")
            # 如果 PyMuPDF 失败，尝试 pypdf
            if self.use_pymupdf and self.use_pypdf:
                logger.info("尝试使用 pypdf 重新加载...")
                try:
                    return self._load_with_pypdf(file_path)
                except Exception as e2:
                    logger.error(f"pypdf 也失败了: {e2}")
                    raise
            raise
    
    def _load_with_pymupdf(self, file_path: Path) -> List[Document]:
        """使用 PyMuPDF 加载 PDF"""
        import pymupdf
        
        logger.debug(f"使用 PyMuPDF 加载: {file_path}")
        
        doc = pymupdf.open(file_path)
        texts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            
            if text.strip():
                texts.append({
                    'text': text,
                    'page_num': page_num + 1
                })
        
        doc.close()
        
        # 创建 Document 对象
        documents = []
        for item in texts:
            metadata = {
                'file_name': file_path.name,
                'file_path': str(file_path),
                'file_type': 'pdf',
                'page_number': item['page_num'],
                'total_pages': len(texts),
            }
            
            doc = Document(
                text=item['text'],
                metadata=metadata
            )
            documents.append(doc)
        
        logger.debug(f"PyMuPDF 加载完成: {len(documents)} 页")
        return documents
    
    def _load_with_pypdf(self, file_path: Path) -> List[Document]:
        """使用 pypdf 加载 PDF"""
        import pypdf
        
        logger.debug(f"使用 pypdf 加载: {file_path}")
        
        documents = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = pypdf.PdfReader(file)
            total_pages = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                
                if text.strip():
                    metadata = {
                        'file_name': file_path.name,
                        'file_path': str(file_path),
                        'file_type': 'pdf',
                        'page_number': page_num,
                        'total_pages': total_pages,
                    }
                    
                    doc = Document(
                        text=text,
                        metadata=metadata
                    )
                    documents.append(doc)
        
        logger.debug(f"pypdf 加载完成: {len(documents)} 页")
        return documents


class DOCXReader:
    """DOCX 文档读取器"""
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("请安装 python-docx: pip install python-docx")
    
    def load_data(self, file_path: Path) -> List[Document]:
        """
        加载 DOCX 文件
        
        Args:
            file_path: DOCX 文件路径
            
        Returns:
            Document 列表
        """
        logger.debug(f"加载 DOCX: {file_path}")
        
        try:
            doc = DocxDocument(file_path)
            
            # 提取所有段落
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # 合并为一个文档
            text = '\n\n'.join(paragraphs)
            
            metadata = {
                'file_name': file_path.name,
                'file_path': str(file_path),
                'file_type': 'docx',
                'paragraph_count': len(paragraphs),
            }
            
            document = Document(
                text=text,
                metadata=metadata
            )
            
            logger.debug(f"DOCX 加载完成: {len(paragraphs)} 段落")
            return [document]
            
        except Exception as e:
            logger.error(f"加载 DOCX 失败 {file_path}: {e}")
            raise


class MarkdownReader:
    """Markdown 文档读取器"""
    
    def load_data(self, file_path: Path) -> List[Document]:
        """
        加载 Markdown 文件
        
        Args:
            file_path: Markdown 文件路径
            
        Returns:
            Document 列表
        """
        logger.debug(f"加载 Markdown: {file_path}")
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            metadata = {
                'file_name': file_path.name,
                'file_path': str(file_path),
                'file_type': 'markdown',
            }
            
            document = Document(
                text=text,
                metadata=metadata
            )
            
            logger.debug(f"Markdown 加载完成: {len(text)} 字符")
            return [document]
            
        except Exception as e:
            logger.error(f"加载 Markdown 失败 {file_path}: {e}")
            raise


class TextCleaner:
    """文本清洗工具"""
    
    @staticmethod
    def clean_text(text: str, preserve_formatting: bool = True) -> str:
        """
        清洗文本
        
        Args:
            text: 原始文本
            preserve_formatting: 是否保留格式信息
            
        Returns:
            清洗后的文本
        """
        if not text:
            return ""
        
        # 移除控制字符（保留换行和制表符）
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        
        # 统一换行符
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        if preserve_formatting:
            # 保留段落结构，移除过多的连续空行
            text = re.sub(r'\n{3,}', '\n\n', text)
            # 移除每行首尾空白，但保留段落缩进
            lines = [line.rstrip() for line in text.split('\n')]
            text = '\n'.join(lines)
        else:
            # 激进清洗：移除所有多余空白
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
        
        # 移除重复空格
        text = re.sub(r' {2,}', ' ', text)
        
        return text
    
    @staticmethod
    def normalize_whitespace(text: str) -> str:
        """规范化空白字符"""
        # 统一中文标点后的空格
        text = re.sub(r'([，。！？；：])\s+', r'\1', text)
        # 统一英文标点后的空格
        text = re.sub(r'([,\.!?;:])\s{2,}', r'\1 ', text)
        return text


class DocumentLoader:
    """
    文档加载器
    
    支持 PDF、DOCX、Markdown 等格式的文档加载
    """
    
    # 支持的文件扩展名
    SUPPORTED_EXTENSIONS = {
        '.pdf': PDFReader,
        '.docx': DOCXReader,
        '.doc': DOCXReader,
        '.md': MarkdownReader,
        '.markdown': MarkdownReader,
    }
    
    def __init__(
        self,
        input_dir: Optional[Union[str, Path]] = None,
        recursive: bool = True,
        clean_text: bool = True,
        preserve_formatting: bool = True,
    ):
        """
        初始化文档加载器
        
        Args:
            input_dir: 输入目录路径
            recursive: 是否递归加载子目录
            clean_text: 是否清洗文本
            preserve_formatting: 是否保留格式信息
        """
        self.input_dir = Path(input_dir) if input_dir else None
        self.recursive = recursive
        self.clean_text = clean_text
        self.preserve_formatting = preserve_formatting
        
        self.text_cleaner = TextCleaner()
        
        # 初始化读取器
        self.readers = {}
        self._initialize_readers()
        
        logger.info(
            f"文档加载器已初始化: "
            f"recursive={recursive}, clean_text={clean_text}, "
            f"preserve_formatting={preserve_formatting}"
        )
    
    def _initialize_readers(self):
        """初始化各类文档读取器"""
        try:
            self.readers['.pdf'] = PDFReader()
            logger.debug("PDF 读取器初始化成功")
        except ImportError as e:
            logger.warning(f"PDF 读取器初始化失败: {e}")
        
        try:
            self.readers['.docx'] = DOCXReader()
            self.readers['.doc'] = self.readers['.docx']
            logger.debug("DOCX 读取器初始化成功")
        except ImportError as e:
            logger.warning(f"DOCX 读取器初始化失败: {e}")
        
        self.readers['.md'] = MarkdownReader()
        self.readers['.markdown'] = self.readers['.md']
        logger.debug("Markdown 读取器初始化成功")
    
    def load_documents(
        self,
        input_dir: Optional[Union[str, Path]] = None,
        file_extensions: Optional[List[str]] = None,
    ) -> List[Document]:
        """
        加载目录下所有支持的文档
        
        Args:
            input_dir: 输入目录（覆盖初始化时的目录）
            file_extensions: 指定要加载的文件扩展名列表
            
        Returns:
            Document 列表
        """
        input_dir = Path(input_dir) if input_dir else self.input_dir
        
        if not input_dir:
            raise ValueError("必须指定 input_dir")
        
        if not input_dir.exists():
            raise FileNotFoundError(f"目录不存在: {input_dir}")
        
        logger.info(f"开始加载文档目录: {input_dir}")
        
        # 获取所有文件
        all_files = self._get_files(input_dir, file_extensions)
        
        if not all_files:
            logger.warning(f"未找到任何支持的文档文件: {input_dir}")
            return []
        
        logger.info(f"找到 {len(all_files)} 个文件")
        
        # 加载所有文档
        all_documents = []
        success_count = 0
        fail_count = 0
        
        for file_path in all_files:
            try:
                documents = self.load_single_document(file_path)
                
                if documents:
                    all_documents.extend(documents)
                    success_count += 1
                    logger.debug(f"✓ 成功加载: {file_path.name} ({len(documents)} 个文档)")
                
            except Exception as e:
                fail_count += 1
                logger.error(f"✗ 加载失败: {file_path.name} - {e}")
        
        logger.info(
            f"文档加载完成: 成功 {success_count}/{len(all_files)}, "
            f"失败 {fail_count}, 总文档数 {len(all_documents)}"
        )
        
        return all_documents
    
    def load_single_document(self, file_path: Union[str, Path]) -> List[Document]:
        """
        加载单个文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            Document 列表
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        # 获取文件扩展名
        ext = file_path.suffix.lower()
        
        if ext not in self.readers:
            raise ValueError(f"不支持的文件格式: {ext}")
        
        # 使用对应的读取器加载
        reader = self.readers[ext]
        documents = reader.load_data(file_path)
        
        # 处理每个文档：清洗文本和添加元数据
        processed_documents = []
        for doc in documents:
            # 获取原始文本
            text = doc.text
            
            # 清洗文本
            if self.clean_text:
                text = self.text_cleaner.clean_text(
                    text,
                    preserve_formatting=self.preserve_formatting
                )
                text = self.text_cleaner.normalize_whitespace(text)
            
            # 创建新文档（合并元数据）
            metadata = doc.metadata.copy() if hasattr(doc, 'metadata') and doc.metadata else {}
            self._add_file_metadata_dict(metadata, file_path)
            
            # 添加文本统计信息
            metadata.update({
                'text_length': len(text),
                'char_count': len(text),
                'word_count': len(text.split()),
            })
            
            processed_doc = Document(
                text=text,
                metadata=metadata,
                id_=doc.id_ if hasattr(doc, 'id_') else None,
            )
            processed_documents.append(processed_doc)
        
        return processed_documents
    
    def _get_files(
        self,
        directory: Path,
        file_extensions: Optional[List[str]] = None,
    ) -> List[Path]:
        """
        获取目录下所有支持的文件
        
        Args:
            directory: 目录路径
            file_extensions: 指定的文件扩展名
            
        Returns:
            文件路径列表
        """
        if file_extensions is None:
            file_extensions = list(self.readers.keys())
        
        # 确保扩展名以点开头
        file_extensions = [
            ext if ext.startswith('.') else f'.{ext}'
            for ext in file_extensions
        ]
        
        files = []
        
        if self.recursive:
            # 递归遍历
            for ext in file_extensions:
                files.extend(directory.glob(f'**/*{ext}'))
        else:
            # 只遍历当前目录
            for ext in file_extensions:
                files.extend(directory.glob(f'*{ext}'))
        
        # 排序以保证一致性
        files = sorted(files)
        
        return files
    
    def _add_file_metadata(self, document: Document, file_path: Path):
        """
        添加文件元数据到 Document 对象
        
        Args:
            document: Document 对象
            file_path: 文件路径
        """
        metadata_dict = {}
        self._add_file_metadata_dict(metadata_dict, file_path)
        document.metadata.update(metadata_dict)
    
    def _add_file_metadata_dict(self, metadata: dict, file_path: Path):
        """
        添加文件元数据到字典
        
        Args:
            metadata: 元数据字典
            file_path: 文件路径
        """
        # 获取文件信息
        stat = file_path.stat()
        
        # 添加元数据
        metadata.update({
            'file_size': stat.st_size,
            'file_size_mb': round(stat.st_size / (1024 * 1024), 2),
            'created_time': datetime.fromtimestamp(stat.st_ctime).isoformat(),
            'modified_time': datetime.fromtimestamp(stat.st_mtime).isoformat(),
        })
    
    def get_document_stats(self, documents: List[Document]) -> Dict[str, Any]:
        """
        获取文档统计信息
        
        Args:
            documents: Document 列表
            
        Returns:
            统计信息字典
        """
        if not documents:
            return {
                'total_documents': 0,
                'total_files': 0,
                'total_size_mb': 0,
                'total_chars': 0,
                'total_words': 0,
                'file_types': {},
                'avg_chars_per_doc': 0,
                'avg_words_per_doc': 0,
            }
        
        # 统计文件类型
        file_types = {}
        file_names = set()
        total_size = 0
        total_chars = 0
        total_words = 0
        
        for doc in documents:
            # 文件类型
            file_type = doc.metadata.get('file_type', 'unknown')
            file_types[file_type] = file_types.get(file_type, 0) + 1
            
            # 文件名（去重）
            file_name = doc.metadata.get('file_name', '')
            if file_name:
                file_names.add(file_name)
            
            # 大小
            total_size += doc.metadata.get('file_size', 0)
            
            # 字符和单词数
            total_chars += doc.metadata.get('char_count', 0)
            total_words += doc.metadata.get('word_count', 0)
        
        stats = {
            'total_documents': len(documents),
            'total_files': len(file_names),
            'total_size_mb': round(total_size / (1024 * 1024), 2),
            'total_chars': total_chars,
            'total_words': total_words,
            'file_types': file_types,
            'avg_chars_per_doc': round(total_chars / len(documents), 2),
            'avg_words_per_doc': round(total_words / len(documents), 2),
        }
        
        return stats
    
    def print_stats(self, documents: List[Document]):
        """打印文档统计信息"""
        stats = self.get_document_stats(documents)
        
        print("\n" + "=" * 60)
        print("文档统计信息")
        print("=" * 60)
        print(f"总文档数: {stats['total_documents']}")
        print(f"总文件数: {stats['total_files']}")
        print(f"总大小: {stats['total_size_mb']} MB")
        print(f"总字符数: {stats['total_chars']:,}")
        print(f"总单词数: {stats['total_words']:,}")
        print(f"\n平均每文档字符数: {stats['avg_chars_per_doc']:,.0f}")
        print(f"平均每文档单词数: {stats['avg_words_per_doc']:,.0f}")
        
        print("\n文件类型分布:")
        for file_type, count in stats['file_types'].items():
            print(f"  {file_type}: {count}")
        
        print("=" * 60 + "\n")


__all__ = [
    'DocumentLoader',
    'PDFReader',
    'DOCXReader',
    'MarkdownReader',
    'TextCleaner',
]
